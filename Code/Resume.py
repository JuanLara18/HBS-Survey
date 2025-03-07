import os
import re
import json
import csv
import time
import logging
import glob

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from dotenv import load_dotenv
import openai

from fpdf import FPDF
from PIL import Image
import docx
import openpyxl
import markdown
import nbformat

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image as ReportLabImage, 
    Table, TableStyle, PageBreak, HRFlowable
)
from reportlab.pdfgen import canvas


# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.FileHandler('report_generator.log'), logging.StreamHandler()]
)
logger = logging.getLogger()

class ProjectReportGenerator:
    def __init__(self, project_root='.'):
        """
        Initialize the report generator with the project root directory and improved context awareness
        """
        self.project_root = project_root
        self.output_pdf = os.path.join(project_root, "Project_Summary_Report.pdf")
        self.temp_dir = os.path.join(project_root, "temp_report_assets")
        self.api_key = None
        self.client = None
        self.doc = None
        self.styles = None
        self.elements = []
        self.image_counter = 0
        
        # Initialize project context information
        self.project_context = {
            "data_files": [],
            "key_analyses": [],
            "visualization_files": [],
            "model_files": [],
            "key_terms": set(),
            "identified_topics": set()
        }
        
        # Create temp directory if it doesn't exist
        if not os.path.exists(self.temp_dir):
            os.makedirs(self.temp_dir)
            
        # Load OpenAI API key
        self._setup_openai()

    def build_project_context(self):
        """Build context about the project by scanning files before detailed analysis"""
        logger.info("Building project context...")
        
        # Scan the project structure to gather context
        for root, dirs, files in os.walk(self.project_root):
            # Skip hidden directories and __pycache__
            dirs[:] = [d for d in dirs if not d.startswith('.') and d != '__pycache__']
            
            for file in files:
                file_path = os.path.join(root, file)
                file_ext = os.path.splitext(file)[1].lower()
                
                # Categorize files
                if file_ext in ['.csv', '.xlsx', '.xls']:
                    self.project_context["data_files"].append(file_path)
                elif file_ext in ['.png', '.jpg', '.jpeg', '.gif']:
                    self.project_context["visualization_files"].append(file_path)
                elif file_ext == '.py' and any(term in file.lower() for term in ['model', 'train', 'predict', 'cluster']):
                    self.project_context["model_files"].append(file_path)
                elif file_ext == '.ipynb':
                    # Try to determine the notebook's purpose
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read().lower()
                            if 'cluster' in content:
                                self.project_context["key_analyses"].append("clustering")
                                self.project_context["identified_topics"].add("clustering")
                            if 'feature_importance' in content or 'feature importance' in content:
                                self.project_context["key_analyses"].append("feature importance")
                                self.project_context["identified_topics"].add("feature importance")
                            if 'regression' in content:
                                self.project_context["key_analyses"].append("regression")
                                self.project_context["identified_topics"].add("regression")
                            if 'classification' in content:
                                self.project_context["key_analyses"].append("classification")
                                self.project_context["identified_topics"].add("classification")
                    except Exception as e:
                        logger.warning(f"Could not scan notebook {file_path}: {str(e)}")
        
        # Extract key terms from filenames
        for file_path in self.project_context["data_files"] + self.project_context["model_files"]:
            filename = os.path.basename(file_path).lower()
            # Extract potential terms from snake_case or camelCase filenames
            terms = re.findall(r'[a-z]+', filename)
            for term in terms:
                if len(term) > 3 and term not in ['data', 'file', 'test', 'train']:
                    self.project_context["key_terms"].add(term)
        
        # Remove duplicates
        self.project_context["key_analyses"] = list(set(self.project_context["key_analyses"]))
        
        logger.info(f"Project context built. Identified topics: {', '.join(self.project_context['identified_topics'])}")

    def _setup_openai(self):
        """Set up OpenAI API client"""
        try:
            # Try to load from .env file first
            load_dotenv(os.path.join(self.project_root, 'Code', 'OPENAI_API_KEY.env'))
            self.api_key = os.getenv('OPENAI_API_KEY')
            
            if not self.api_key:
                # If not found, try loading from a different location or format
                env_path = os.path.join(self.project_root, 'Code', 'OPENAI_API_KEY.env')
                if os.path.exists(env_path):
                    with open(env_path, 'r') as f:
                        content = f.read().strip()
                        if content.startswith('OPENAI_API_KEY='):
                            self.api_key = content.split('=')[1].strip('"\'')
                        else:
                            self.api_key = content
            
            if self.api_key:
                openai.api_key = self.api_key
                self.client = openai.OpenAI(api_key=self.api_key)
                logger.info("OpenAI API client initialized successfully")
            else:
                logger.warning("OpenAI API key not found. ChatGPT features will be disabled.")
        except Exception as e:
            logger.error(f"Error setting up OpenAI client: {str(e)}")
            self.client = None

    def _query_chatgpt(self, prompt, max_retries=3, retry_delay=5):
        """
        Query ChatGPT with error handling and retries
        """
        if not self.client:
            return "ChatGPT integration not available (API key not found)."
        
        for attempt in range(max_retries):
            try:
                response = self.client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are a helpful data science assistant. Provide clear, concise explanations of code and analytical results."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=1000
                )
                return response.choices[0].message.content
            except Exception as e:
                logger.warning(f"ChatGPT query failed (attempt {attempt+1}/{max_retries}): {str(e)}")
                if attempt < max_retries - 1:
                    time.sleep(retry_delay)
                else:
                    logger.error("All ChatGPT query attempts failed")
                    return f"Error querying ChatGPT: {str(e)}"
    
    def initialize_document(self):
        """Initialize the PDF document with improved styling"""
        # Configure page and margins
        self.doc = SimpleDocTemplate(
            self.output_pdf,
            pagesize=A4,
            rightMargin=54,  # Reduced margins for better space usage
            leftMargin=54,
            topMargin=72,
            bottomMargin=54
        )
        
        # Get base styles and improve them
        self.styles = getSampleStyleSheet()
        
        # Improve title style
        self.styles['Title'].fontSize = 20
        self.styles['Title'].leading = 24
        self.styles['Title'].spaceBefore = 12
        self.styles['Title'].spaceAfter = 12
        self.styles['Title'].alignment = 1  # Centered
        self.styles['Title'].textColor = colors.darkblue
        
        # Improve heading styles
        self.styles['Heading1'].fontSize = 18
        self.styles['Heading1'].leading = 22
        self.styles['Heading1'].spaceBefore = 24
        self.styles['Heading1'].spaceAfter = 12
        self.styles['Heading1'].textColor = colors.darkblue
        self.styles['Heading1'].borderWidth = 0
        self.styles['Heading1'].borderPadding = 5
        self.styles['Heading1'].borderColor = colors.lightgrey
        self.styles['Heading1'].borderRadius = 2
        
        self.styles['Heading2'].fontSize = 14
        self.styles['Heading2'].leading = 18
        self.styles['Heading2'].spaceBefore = 18
        self.styles['Heading2'].spaceAfter = 12
        self.styles['Heading2'].textColor = colors.navy
        
        self.styles['Heading3'].fontSize = 12
        self.styles['Heading3'].leading = 14
        self.styles['Heading3'].spaceBefore = 14
        self.styles['Heading3'].spaceAfter = 8
        self.styles['Heading3'].textColor = colors.darkblue
        
        # Improve normal style for text
        self.styles['Normal'].fontSize = 10
        self.styles['Normal'].leading = 14
        self.styles['Normal'].spaceBefore = 6
        self.styles['Normal'].spaceAfter = 10
        
        # Improved code style
        if 'Code' not in self.styles:
            self.styles.add(ParagraphStyle(
                name='Code',
                parent=self.styles['Normal'],
                fontName='Courier',
                fontSize=9,
                leading=12,
                firstLineIndent=0,
                leftIndent=20,
                rightIndent=20,
                backColor=colors.lightgrey,
                borderColor=colors.grey,
                borderWidth=1,
                borderPadding=5,
                borderRadius=2,
                spaceBefore=10,
                spaceAfter=10
            ))
        else:
            # Update existing style
            self.styles['Code'].fontName = 'Courier'
            self.styles['Code'].fontSize = 9
            self.styles['Code'].leading = 12
            self.styles['Code'].firstLineIndent = 0
            self.styles['Code'].leftIndent = 20
            self.styles['Code'].rightIndent = 20
            self.styles['Code'].backColor = colors.lightgrey
            self.styles['Code'].borderColor = colors.grey
            self.styles['Code'].borderWidth = 1
            self.styles['Code'].borderPadding = 5
            self.styles['Code'].borderRadius = 2
            self.styles['Code'].spaceBefore = 10
            self.styles['Code'].spaceAfter = 10
        
        # Style for captions
        self.styles.add(ParagraphStyle(
            name='Caption',
            parent=self.styles['Normal'],
            fontSize=9,
            leading=12,
            alignment=1,  # Centered
            italics=True,
            textColor=colors.darkslategray,
            spaceBefore=2,
            spaceAfter=12
        ))
        
        # Style for list items
        self.styles.add(ParagraphStyle(
            name='ListItem',
            parent=self.styles['Normal'],
            fontSize=10,
            leading=14,
            leftIndent=20,
            firstLineIndent=-15,
            bulletIndent=0,
            spaceBefore=2,
            spaceAfter=2
        ))
        
        # Clear elements list
        self.elements = []
        
        # Add cover page
        self.add_cover_page()
        
        # Add page break
        self.elements.append(PageBreak())
        
        # Add table of contents title
        self.elements.append(Paragraph("Table of Contents", self.styles['Heading1']))
        self.elements.append(Spacer(1, 0.3*inch))
        
        # Here we would add the table of contents (will be generated later)
        # For now add a placeholder
        self.elements.append(Paragraph("<i>The table of contents will be automatically generated.</i>", self.styles['Normal']))
        self.elements.append(Spacer(1, 0.5*inch))
        
        # Add page break to start content
        self.elements.append(PageBreak())

    def add_cover_page(self):
        """Add an elegant cover page to the document"""
        # Main title
        title_style = ParagraphStyle(
            name='CoverTitle',
            parent=self.styles['Title'],
            fontSize=24,
            leading=30,
            alignment=1,  # Centered
            textColor=colors.darkblue,
            spaceBefore=100  # Push down to center on page
        )
        
        self.elements.append(Paragraph("HBS Survey Project", title_style))
        self.elements.append(Spacer(1, 0.3*inch))
        
        subtitle_style = ParagraphStyle(
            name='CoverSubtitle',
            parent=self.styles['Title'],
            fontSize=18,
            leading=22,
            alignment=1,  # Centered
            textColor=colors.navy
        )
        
        self.elements.append(Paragraph("Comprehensive Analysis Report", subtitle_style))
        self.elements.append(Spacer(1, 2*inch))
        
        # Generation information
        info_style = ParagraphStyle(
            name='CoverInfo',
            parent=self.styles['Normal'],
            fontSize=12,
            leading=16,
            alignment=1  # Centered
        )
        
        date_text = f"Generated on: {time.strftime('%B %d, %Y, %H:%M:%S')}"
        self.elements.append(Paragraph(date_text, info_style))
        self.elements.append(Spacer(1, 0.5*inch))
        
        # Add footer note
        footer_style = ParagraphStyle(
            name='CoverFooter',
            parent=self.styles['Normal'],
            fontSize=10,
            leading=12,
            alignment=1,  # Centered
            textColor=colors.darkgrey
        )
        
        self.elements.append(Spacer(1, 3*inch))
        self.elements.append(Paragraph("Research Project Based on Data Analysis", footer_style))

    def add_heading(self, text, level=1):
        """Add a heading to the document with improved formatting"""
        style_name = f"Heading{level}"
        
        # Add additional space before level 1 headings
        if level == 1:
            self.elements.append(Spacer(1, 0.3*inch))
        else:
            self.elements.append(Spacer(1, 0.1*inch))
        
        # Add horizontal line before level 1 headings
        if level == 1:
            self.elements.append(HRFlowable(
                width="100%",
                thickness=1,
                color=colors.lightgrey,
                spaceBefore=10,
                spaceAfter=10
            ))
        
        self.elements.append(Paragraph(text, self.styles[style_name]))
    
    def add_paragraph(self, text, style_name='Normal'):
        """Add a paragraph to the document with specified style"""
        self.elements.append(Paragraph(text, self.styles[style_name]))
                     
    def add_bullet_point(self, text):
        """Add a bullet point with proper formatting"""
        bullet_text = f"• {text}"
        self.elements.append(Paragraph(bullet_text, self.styles['ListItem']))                         
                         
    def add_code(self, code):
        """Add code with improved formatting"""
        # Encapsulate in a single paragraph for better appearance
        formatted_code = code.replace('\n', '<br/>')
        formatted_code = formatted_code.replace(' ', '&nbsp;')  # Preserve spaces
        code_paragraph = Paragraph(formatted_code, self.styles['Code'])
        self.elements.append(code_paragraph)

    def add_image(self, image_path, width=6*inch, caption=None):
        """Add an image to the document with proper sizing and optional caption"""
        try:
            # Check if image exists
            if not os.path.exists(image_path):
                logger.warning(f"Image not found: {image_path}")
                return False
            
            # Resize large images
            try:
                from PIL import Image as PILImage
                img_pil = PILImage.open(image_path)
                img_width, img_height = img_pil.size
                
                # Calculate aspect ratio
                aspect_ratio = img_height / img_width
                
                # Limit maximum width to 5 inches for large images
                max_width = 5 * inch
                if width > max_width:
                    width = max_width
                
                # Calculate proportional height
                height = width * aspect_ratio
                
                # Limit maximum height to 6 inches
                max_height = 6 * inch
                if height > max_height:
                    height = max_height
                    width = height / aspect_ratio
            except Exception as e:
                logger.warning(f"Could not determine image dimensions for {image_path}: {str(e)}")
                # Default safe values
                width = 4 * inch
                height = None
            
            # Center the image on the page
            img_container = Table([[ReportLabImage(image_path, width=width, height=height)]], 
                                colWidths=[self.doc.width])
            img_container.setStyle(TableStyle([('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                                            ('LEFTPADDING', (0, 0), (-1, -1), 0),
                                            ('RIGHTPADDING', (0, 0), (-1, -1), 0)]))
            self.elements.append(img_container)
            
            # Add caption if provided
            if caption:
                self.elements.append(Paragraph(caption, self.styles['Caption']))
            
            self.elements.append(Spacer(1, 0.2*inch))
            return True
        except Exception as e:
            logger.error(f"Failed to add image {image_path}: {str(e)}")
            return False

    def add_table(self, data, col_widths=None, highlight_header=True, alternating_colors=True):
        """Add a table with improved formatting"""
        try:
            if not data:
                return
            
            # Create table
            table = Table(data, colWidths=col_widths)
            
            # Base styles
            table_style = [
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
                ('BOX', (0, 0), (-1, -1), 1, colors.darkgrey),
            ]
            
            # Highlight header
            if highlight_header:
                table_style.append(('BACKGROUND', (0, 0), (-1, 0), colors.lightblue))
                table_style.append(('TEXTCOLOR', (0, 0), (-1, 0), colors.darkblue))
            
            # Alternating colors for rows
            if alternating_colors and len(data) > 1:
                for i in range(1, len(data), 2):
                    table_style.append(('BACKGROUND', (0, i), (-1, i), colors.lightgrey))
            
            table.setStyle(TableStyle(table_style))
            
            # Use a container to center the table
            table_container = Table([[table]], colWidths=[self.doc.width])
            table_container.setStyle(TableStyle([('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                                                ('RIGHTPADDING', (0, 0), (-1, -1), 0)]))
            
            self.elements.append(table_container)
            self.elements.append(Spacer(1, 0.25*inch))
        except Exception as e:
            logger.error(f"Failed to add table: {str(e)}")
    
    def process_code_file(self, file_path):
        """Process a single code file and extract insights"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                code_content = f.read()
            
            file_name = os.path.basename(file_path)
            file_ext = os.path.splitext(file_name)[1].lower()
            
            self.add_heading(f"File: {file_name}", level=2)
            
            # For Python files, generate a summary using ChatGPT
            if file_ext == '.py':
                prompt = f"""
                Analyze this Python code and provide a concise summary (maximum 200 words) of:
                1. What the code does
                2. Key functions/classes
                3. Any important algorithms or techniques used
                
                CODE:
                {code_content[:4000]}  # Limit to avoid token limits
                """
                
                summary = self._query_chatgpt(prompt)
                self.add_paragraph("Code Summary:")
                self.add_paragraph(summary)
                
                # Add shortened code sample (first 30 lines)
                code_lines = code_content.split('\n')
                sample = '\n'.join(code_lines[:min(30, len(code_lines))])
                
                self.add_paragraph("Code Sample (first 30 lines):")
                self.add_code(sample)
                
                if len(code_lines) > 30:
                    self.add_paragraph(f"(... {len(code_lines) - 30} more lines not shown ...)")
                
            # For Jupyter notebooks, handle differently
            elif file_ext == '.ipynb':
                self.process_notebook(file_path)
                
            else:
                # For other file types, just add a brief mention
                self.add_paragraph(f"File type: {file_ext} - Content size: {len(code_content)} characters")
                
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            self.add_paragraph(f"Error processing this file: {str(e)}")

    def process_notebook(self, notebook_path):
        """Process a Jupyter notebook with improved content summarization"""
        try:
            with open(notebook_path, 'r', encoding='utf-8') as f:
                notebook_content = json.load(f)
            
            # Check basic structure
            if "cells" not in notebook_content:
                self.add_paragraph(f"Notebook does not have a valid format (no 'cells' found)")
                return
                    
            cells = notebook_content.get("cells", [])
            
            # Filter cells by type
            markdown_cells = []
            code_cells = []
            
            for c in cells:
                if "cell_type" not in c:
                    continue
                    
                if c["cell_type"] == "markdown" and "source" in c:
                    source = c.get("source", "")
                    if isinstance(source, list):
                        source = "".join(source)
                    markdown_cells.append(source)
                elif c["cell_type"] == "code" and "source" in c:
                    source = c.get("source", "")
                    if isinstance(source, list):
                        source = "".join(source)
                    code_cells.append(source)
            
            file_name = os.path.basename(notebook_path)
            self.add_heading(f"Notebook Analysis: {file_name}", level=2)
            self.add_paragraph(f"Notebook contains {len(cells)} cells ({len(markdown_cells)} markdown, {len(code_cells)} code)")
            
            # Extract key information for better summarization
            notebook_info = {
                "imports": [],
                "functions": [],
                "visualizations": [],
                "model_types": [],
                "data_operations": []
            }
            
            # Look for imports, functions, and visualizations in code
            import_pattern = re.compile(r'import\s+(\w+)|from\s+(\w+)\s+import')
            viz_patterns = ["plt.", "sns.", ".plot(", ".imshow(", ".figure", "px.", ".scatter(", ".bar(", ".hist("]
            model_patterns = ["LinearRegression", "RandomForest", "LogisticRegression", "KMeans", "DBSCAN", "cluster", 
                            "SVC", "DecisionTree", "XGBoost", "model.fit", "train_test_split"]
            data_patterns = ["pd.read_", "DataFrame", ".groupby", ".pivot", ".merge", ".join", ".concat", ".value_counts()"]
            
            for code in code_cells:
                # Find imports
                imports = import_pattern.findall(code)
                for imp in imports:
                    if imp[0]:  # import x
                        notebook_info["imports"].append(imp[0])
                    elif imp[1]:  # from x import
                        notebook_info["imports"].append(imp[1])
                
                # Look for function definitions
                function_matches = re.findall(r'def\s+(\w+)\s*\(', code)
                notebook_info["functions"].extend(function_matches)
                
                # Check for visualizations
                for pattern in viz_patterns:
                    if pattern in code:
                        notebook_info["visualizations"].append(pattern.strip(".()"))
                        
                # Check for ML models
                for pattern in model_patterns:
                    if pattern in code:
                        notebook_info["model_types"].append(pattern)
                        
                # Check for data operations
                for pattern in data_patterns:
                    if pattern in code:
                        notebook_info["data_operations"].append(pattern)
            
            # Remove duplicates and limit entries
            for key in notebook_info:
                notebook_info[key] = list(set(notebook_info[key]))[:5]  # Take only the first 5 unique items
            
            # Extract headings from markdown for better structure understanding
            headings = []
            heading_pattern = re.compile(r'^(#+)\s+(.*?)$', re.MULTILINE)
            
            for md in markdown_cells:
                for match in heading_pattern.finditer(md):
                    level = len(match.group(1))
                    heading_text = match.group(2).strip()
                    headings.append((level, heading_text))
            
            # Sort headings by their appearance (assuming they are in order in the notebook)
            headings = sorted(headings, key=lambda x: (x[0], headings.index((x[0], x[1]))))
            
            # Display notebook structure
            if headings:
                self.add_heading("Notebook Structure", level=3)
                for level, heading in headings[:10]:  # Limit to first 10 headings
                    indent = "  " * (level - 1)
                    self.add_paragraph(f"{indent}• {heading}")
                
                if len(headings) > 10:
                    self.add_paragraph(f"(... {len(headings) - 10} more headings not shown ...)")
            
            # Generate a smarter summary using extracted info
            if any(len(v) > 0 for v in notebook_info.values()):
                self.add_heading("Notebook Technical Content", level=3)
                
                if notebook_info["imports"]:
                    libraries = ", ".join(notebook_info["imports"])
                    self.add_paragraph(f"Key libraries: {libraries}")
                
                if notebook_info["functions"]:
                    functions = ", ".join(notebook_info["functions"])
                    self.add_paragraph(f"Custom functions: {functions}")
                
                if notebook_info["visualizations"]:
                    visualizations = ", ".join(set(notebook_info["visualizations"]))
                    self.add_paragraph(f"Visualization methods: {visualizations}")
                    
                if notebook_info["model_types"]:
                    models = ", ".join(set(notebook_info["model_types"]))
                    self.add_paragraph(f"Analysis techniques: {models}")
                    
                if notebook_info["data_operations"]:
                    operations = ", ".join(set(notebook_info["data_operations"]))
                    self.add_paragraph(f"Data operations: {operations}")
                    
            # Combine markdown text and code snippets for better context
            md_sample = "\n".join(markdown_cells[:3])
            code_sample = "\n".join(code_cells[:3])
            
            # Create a more targeted prompt based on what we've found
            analysis_type = ""
            if any("cluster" in model.lower() for model in notebook_info["model_types"]):
                analysis_type = "clustering analysis"
            elif any("regress" in model.lower() for model in notebook_info["model_types"]):
                analysis_type = "regression analysis"
            elif any("classif" in model.lower() for model in notebook_info["model_types"]):
                analysis_type = "classification analysis"
            elif any("feature" in func.lower() or "importance" in func.lower() for func in notebook_info["functions"]):
                analysis_type = "feature importance analysis"
            
            prompt = f"""
            Analyze this Jupyter notebook content and provide a concise summary (maximum 250 words) of:
            1. The main purpose/topic of the notebook (seems to be {analysis_type if analysis_type else 'data analysis'})
            2. Key analyses or visualizations it contains
            3. The main findings or conclusions (if apparent)
            
            Key libraries used: {', '.join(notebook_info['imports'])}
            
            NOTEBOOK HEADINGS:
            {', '.join([h[1] for h in headings[:5]])}
            
            NOTEBOOK SAMPLE:
            
            Markdown cells:
            {md_sample[:1500]}
            
            Code cells:
            {code_sample[:1500]}
            """
            
            summary = self._query_chatgpt(prompt)
            
            # Clean and display the summary
            self.add_heading("Notebook Summary:", level=3)
            # Split into paragraphs for better readability
            summary_paragraphs = summary.split('\n\n')
            for paragraph in summary_paragraphs:
                if paragraph.strip():
                    # Clean any remaining markdown
                    clean_para = self._clean_markdown(paragraph)
                    self.add_paragraph(clean_para)
            
        except Exception as e:
            logger.error(f"Error processing notebook {notebook_path}: {str(e)}")
            self.add_paragraph(f"Error processing this notebook: {str(e)}")

    def _clean_markdown(self, text):
        """Clean markdown formatting for better display in PDF"""
        if not text:
            return ""
        
        # Replace bold markers
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'__(.*?)__', r'\1', text)
        
        # Replace italic markers
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'_(.*?)_', r'\1', text)
        
        # Handle headings
        text = re.sub(r'^#+\s+(.*?)$', r'\1', text, flags=re.MULTILINE)
        
        # Handle numbered lists
        text = re.sub(r'^\d+\.\s+(.*?)$', r'• \1', text, flags=re.MULTILINE)
        
        # Handle bullet lists
        text = re.sub(r'^\-\s+(.*?)$', r'• \1', text, flags=re.MULTILINE)
        text = re.sub(r'^\*\s+(.*?)$', r'• \1', text, flags=re.MULTILINE)
        
        return text

    def add_formatted_paragraph(self, markdown_text, style_name='Normal'):
        """Add a paragraph with proper handling of markdown formatting"""
        if not markdown_text:
            return
        
        # First clean up basic markdown elements
        clean_text = self._clean_markdown(markdown_text)
        
        # Now add with the selected style
        self.add_paragraph(clean_text, style_name)

    def process_image_file(self, image_path):
        """Process an image file and add it to the report"""
        try:
            file_name = os.path.basename(image_path)
            self.add_heading(f"Image: {file_name}", level=3)
            
            # Add the image to the document
            success = self.add_image(image_path, caption=file_name)
            
            if not success:
                self.add_paragraph(f"Failed to add image: {file_name}")
                
            # For plot images, try to interpret what they show
            if 'plot' in file_name.lower() or 'figure' in file_name.lower() or 'chart' in file_name.lower() or any(x in file_name.lower() for x in ['scatter', 'bar', 'histogram', 'heatmap', 'cluster']):
                prompt = f"""
                This is a data visualization image named "{file_name}" from a data science project.
                Based only on the filename, what might this visualization be showing? 
                Provide a brief, educated guess (2-3 sentences) about what information this plot might be visualizing.
                """
                
                interpretation = self._query_chatgpt(prompt)
                self.add_paragraph("Possible interpretation:")
                self.add_paragraph(interpretation)
                
        except Exception as e:
            logger.error(f"Error processing image {image_path}: {str(e)}")
            self.add_paragraph(f"Error processing this image: {str(e)}")
    
    def process_excel_file(self, excel_path):
        """Process an Excel file and extract key information"""
        try:
            file_name = os.path.basename(excel_path)
            self.add_heading(f"Excel File: {file_name}", level=2)
            
            # Load the workbook
            wb = openpyxl.load_workbook(excel_path, read_only=True, data_only=True)
            sheet_names = wb.sheetnames
            
            self.add_paragraph(f"Contains {len(sheet_names)} sheets: {', '.join(sheet_names)}")
            
            # Process each sheet (limit to first 3 for brevity)
            for sheet_name in sheet_names[:3]:
                sheet = wb[sheet_name]
                self.add_heading(f"Sheet: {sheet_name}", level=3)
                
                # Get dimensions
                if sheet.max_row < 1000 and sheet.max_column < 50:  # Reasonable size check
                    rows = min(6, sheet.max_row)  # First 5 rows plus header
                    cols = min(10, sheet.max_column)  # First 10 columns
                    
                    # Create data for table
                    data = []
                    for r in range(1, rows + 1):
                        row_data = []
                        for c in range(1, cols + 1):
                            cell_value = sheet.cell(row=r, column=c).value
                            row_data.append(str(cell_value) if cell_value is not None else "")
                        data.append(row_data)
                    
                    # Add table to document
                    if data:
                        self.add_paragraph(f"Preview of first {rows} rows and {cols} columns:")
                        self.add_table(data)
                        
                        if sheet.max_column > cols:
                            self.add_paragraph(f"(... {sheet.max_column - cols} more columns not shown ...)")
                        
                        if sheet.max_row > rows:
                            self.add_paragraph(f"(... {sheet.max_row - rows} more rows not shown ...)")
                else:
                    self.add_paragraph(f"Sheet is too large to display: {sheet.max_row} rows x {sheet.max_column} columns")
                    
                    # Try to read just the headers
                    headers = [str(sheet.cell(row=1, column=c).value) for c in range(1, min(15, sheet.max_column) + 1)]
                    headers = [h if h != "None" else f"Column {i}" for i, h in enumerate(headers)]
                    
                    self.add_paragraph("Column headers:")
                    self.add_paragraph(", ".join(headers))
                    
                    if sheet.max_column > 15:
                        self.add_paragraph(f"(... {sheet.max_column - 15} more columns not shown ...)")
            
            if len(sheet_names) > 3:
                self.add_paragraph(f"(... {len(sheet_names) - 3} more sheets not shown ...)")
                
            # Generate a summary interpretation
            prompt = f"""
            This Excel file "{file_name}" has the following sheets: {', '.join(sheet_names)}.
            Based on this information and considering the context of a data analysis project, 
            what insights or data might this file contain? What role might it play in the analysis?
            Please provide a brief hypothesis (3-4 sentences).
            """
            
            interpretation = self._query_chatgpt(prompt)
            self.add_paragraph("Possible content interpretation:")
            self.add_paragraph(interpretation)
            
        except Exception as e:
            logger.error(f"Error processing Excel file {excel_path}: {str(e)}")
            self.add_paragraph(f"Error processing this Excel file: {str(e)}")
    
    def process_csv_file(self, csv_path):
        """Process a CSV file and extract key information"""
        try:
            file_name = os.path.basename(csv_path)
            self.add_heading(f"CSV File: {file_name}", level=2)
            
            # Try to read with pandas for better handling
            try:
                df = pd.read_csv(csv_path, nrows=10)  # Read just 10 rows for preview
                
                # Get basic info
                full_df_info = pd.read_csv(csv_path, nrows=0)  # Just for column info
                num_cols = len(full_df_info.columns)
                
                # Count lines in file for row count (more efficient than loading entire DataFrame)
                with open(csv_path, 'r', encoding='utf-8', errors='ignore') as f:
                    num_rows = sum(1 for _ in f) - 1  # Subtract 1 for header
                
                self.add_paragraph(f"Contains {num_rows} rows and {num_cols} columns")
                
                # Column names
                self.add_paragraph("Columns:")
                self.add_paragraph(", ".join(df.columns.tolist()))
                
                # Data preview
                if not df.empty:
                    # Convert DataFrame to list of lists for the table
                    table_data = [df.columns.tolist()]  # Header row
                    table_data.extend(df.values.tolist())
                    
                    # Convert all elements to strings
                    table_data = [[str(cell) for cell in row] for row in table_data]
                    
                    self.add_paragraph("Data preview (first 10 rows):")
                    self.add_table(table_data)
                
            except Exception as e:
                # Fallback to basic CSV reading if pandas fails
                logger.warning(f"Pandas reading failed for {csv_path}, falling back to CSV reader: {str(e)}")
                
                with open(csv_path, 'r', encoding='utf-8', errors='ignore') as f:
                    csv_reader = csv.reader(f)
                    headers = next(csv_reader)
                    
                    # Count rows
                    row_count = sum(1 for _ in csv_reader)
                    
                self.add_paragraph(f"Contains {row_count} rows and {len(headers)} columns")
                self.add_paragraph("Columns:")
                self.add_paragraph(", ".join(headers))
            
            # Generate a summary interpretation
            prompt = f"""
            This CSV file "{file_name}" appears to contain data with columns: {', '.join(df.columns.tolist() if 'df' in locals() else headers)}.
            Based on the file name and column headers, what kind of data might this contain? 
            What insights could it provide to the overall analysis?
            Please provide a brief hypothesis (3-4 sentences).
            """
            
            interpretation = self._query_chatgpt(prompt)
            self.add_paragraph("Possible data interpretation:")
            self.add_paragraph(interpretation)
            
        except Exception as e:
            logger.error(f"Error processing CSV file {csv_path}: {str(e)}")
            self.add_paragraph(f"Error processing this CSV file: {str(e)}")
    
    def process_docx_file(self, docx_path):
        """Process a Word document and extract key information"""
        try:
            file_name = os.path.basename(docx_path)
            self.add_heading(f"Word Document: {file_name}", level=2)
            
            # Open the document
            doc = docx.Document(docx_path)
            
            # Extract paragraphs and headings
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Find headings (assuming they use heading styles)
            headings = []
            for p in doc.paragraphs:
                if p.style.name.startswith('Heading'):
                    headings.append(p.text)
            
            # Document stats
            self.add_paragraph(f"Document contains {len(paragraphs)} paragraphs and {len(headings)} headings")
            
            # Show document structure if headings exist
            if headings:
                self.add_paragraph("Document structure:")
                for heading in headings[:10]:  # Limit to 10 headings
                    self.add_paragraph(f"- {heading}")
                
                if len(headings) > 10:
                    self.add_paragraph(f"(... {len(headings) - 10} more headings not shown ...)")
            
            # Show document intro (first few paragraphs)
            if paragraphs:
                self.add_paragraph("Document introduction:")
                for p in paragraphs[:5]:  # Show first 5 paragraphs
                    self.add_paragraph(p)
                
                if len(paragraphs) > 5:
                    self.add_paragraph(f"(... {len(paragraphs) - 5} more paragraphs not shown ...)")
            
            # Generate a summary of the document
            doc_sample = "\n".join(paragraphs[:10])
            prompt = f"""
            This document "{file_name}" has the following structure and content:
            
            Headings: {', '.join(headings[:5])}
            
            Sample content:
            {doc_sample[:1500]}
            
            Based on this information, provide a concise summary (maximum 200 words) of what this document appears to contain and its significance to the project.
            """
            
            summary = self._query_chatgpt(prompt)
            self.add_paragraph("Document Summary:")
            self.add_paragraph(summary)
            
        except Exception as e:
            logger.error(f"Error processing Word document {docx_path}: {str(e)}")
            self.add_paragraph(f"Error processing this Word document: {str(e)}")
    
    def process_markdown_file(self, md_path):
        """Process a Markdown file and extract key information"""
        try:
            file_name = os.path.basename(md_path)
            self.add_heading(f"Markdown File: {file_name}", level=2)
            
            # Read the file
            with open(md_path, 'r', encoding='utf-8', errors='ignore') as f:
                md_content = f.read()
            
            # Display the content
            self.add_paragraph("File content:")
            
            # Parse Markdown to extract structure
            lines = md_content.split('\n')
            headings = []
            
            for line in lines:
                line = line.strip()
                if line.startswith('#'):
                    headings.append(line)
            
            # Show headings
            if headings:
                self.add_paragraph("Markdown structure:")
                for heading in headings:
                    self.add_paragraph(f"- {heading}")
            
            # Convert to HTML for cleaner display
            html = markdown.markdown(md_content)
            
            # Show sample content (limit length)
            if len(md_content) > 1000:
                preview = md_content[:1000] + "..."
            else:
                preview = md_content
                
            self.add_paragraph("Content preview:")
            self.add_paragraph(preview)
            
            # Generate a summary
            prompt = f"""
            This Markdown file "{file_name}" contains the following content:
            
            {md_content[:1500]}
            
            Please provide a concise summary (maximum 150 words) of what this document contains and its purpose in the project.
            """
            
            summary = self._query_chatgpt(prompt)
            self.add_paragraph("Document Summary:")
            self.add_paragraph(summary)
            
        except Exception as e:
            logger.error(f"Error processing Markdown file {md_path}: {str(e)}")
            self.add_paragraph(f"Error processing this Markdown file: {str(e)}")
    
    def explore_directory(self, directory_path, section_title=None):
        """
        Recursively explore a directory and process its contents, avoiding already processed files
        """
        if section_title:
            self.add_heading(section_title, level=1)
        
        # Get all files and dirs, sorted alphabetically
        try:
            all_items = sorted(os.listdir(directory_path))
            files = [f for f in all_items if os.path.isfile(os.path.join(directory_path, f))]
            dirs = [d for d in all_items if os.path.isdir(os.path.join(directory_path, d))]
            
            # Skip hidden files and certain directories
            files = [f for f in files if not f.startswith('.') and f != '__pycache__']
            dirs = [d for d in dirs if not d.startswith('.') and d != '__pycache__']
            
            # Add directory summary
            if files or dirs:
                self.add_paragraph(f"Directory contains {len(files)} files and {len(dirs)} subdirectories")
                
                if files:
                    self.add_paragraph("Files:")
                    file_list = ", ".join(files)
                    self.add_paragraph(file_list)
                
                if dirs:
                    self.add_paragraph("Subdirectories:")
                    dir_list = ", ".join(dirs)
                    self.add_paragraph(dir_list)
            
            # Process individual files
            for file in files:
                file_path = os.path.join(directory_path, file)
                
                # Skip already processed files to avoid repetition
                if file_path in self.processed_files:
                    continue
                    
                self.processed_files.add(file_path)
                
                file_ext = os.path.splitext(file)[1].lower()
                
                # Process different file types
                if file_ext in ['.py']:
                    self.process_code_file(file_path)
                elif file_ext == '.ipynb':
                    self.process_notebook(file_path)
                elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                    self.process_image_file(file_path)
                elif file_ext in ['.xlsx', '.xls']:
                    self.process_excel_file(file_path)
                elif file_ext == '.csv':
                    self.process_csv_file(file_path)
                elif file_ext == '.docx':
                    self.process_docx_file(file_path)
                elif file_ext in ['.md', '.markdown']:
                    self.process_markdown_file(file_path)
                else:
                    # Other file types, just add a brief entry
                    file_size = os.path.getsize(file_path) / 1024  # Size in KB
                    self.add_heading(f"File: {file}", level=3)
                    self.add_paragraph(f"Type: {file_ext}, Size: {file_size:.2f} KB")
            
            # Process subdirectories recursively
            for subdir in dirs:
                subdir_path = os.path.join(directory_path, subdir)
                self.add_heading(f"Subdirectory: {subdir}", level=2)
                self.explore_directory(subdir_path)
                
        except Exception as e:
            logger.error(f"Error exploring directory {directory_path}: {str(e)}")
            self.add_paragraph(f"Error exploring this directory: {str(e)}")

    def process_project(self):
        """Process the entire project structure with improved organization and reduced redundancy"""
        try:
            # Initialize the document
            self.initialize_document()
            
            # Keep track of processed files to avoid duplication
            self.processed_files = set()
            
            # Add executive summary
            self.add_heading("Executive Summary", level=1)
            
            # Generate project summary
            project_structure = self.get_project_structure()
            prompt = f"""
            Generate a concise executive summary for a data science project with the following structure:
            
            {project_structure}
            
            Include:
            1. A clear description of the project's main objective
            2. The analytical methods used
            3. Key findings discovered
            4. High-level recommendations
            
            Make it concise (200-250 words) but informative, written in a professional and executive tone.
            """
            
            executive_summary = self._query_chatgpt(prompt)
            self.add_paragraph(executive_summary)
            
            # Add page break after executive summary
            self.elements.append(PageBreak())
            
            # Add project overview
            self.add_heading("Project Overview", level=1)
            
            # Get a detailed overview of the project
            prompt = f"""
            Generate a detailed description for a data science project with the following structure:
            
            {project_structure}
            
            Include:
            1. High-level description of the project's purpose
            2. The main components and their purpose
            3. The analytical workflow evident from the directory structure
            4. What kind of results are present
            
            Make it detailed and insightful (250-300 words), with a professional and technical focus.
            """
            
            overview = self._query_chatgpt(prompt)
            self.add_paragraph(overview)
            
            # Process important files first to avoid duplicating in directory exploration
            self.process_key_files()
            
            # Add page breaks between major sections
            self.elements.append(PageBreak())
            
            # Process code directory
            code_dir = os.path.join(self.project_root, "Code")
            if os.path.exists(code_dir):
                self.explore_directory(code_dir, "Code Analysis")
            
            # Add page break before data section
            self.elements.append(PageBreak())
            
            # Process data files
            data_dir = os.path.join(self.project_root, "Data")
            if os.path.exists(data_dir):
                self.explore_directory(data_dir, "Data Analysis")
            
            # Add page break before results section
            self.elements.append(PageBreak())
            
            # Process output directory
            output_dir = os.path.join(self.project_root, "Output")
            if os.path.exists(output_dir):
                self.explore_directory(output_dir, "Results Analysis")
            
            # Add page break before conclusion
            self.elements.append(PageBreak())
            
            # Generate conclusion and insights
            self.generate_conclusion()
            
            # Save the document
            self.doc.build(self.elements, canvasmaker=PageNumCanvas)
            logger.info(f"Report successfully saved to {self.output_pdf}")
            
            return self.output_pdf
        except Exception as e:
            logger.error(f"Error processing project: {str(e)}")
            # Try to save what we have so far
            if self.elements:
                try:
                    self.doc.build(self.elements, canvasmaker=PageNumCanvas)
                    logger.info(f"Partial report saved to {self.output_pdf}")
                except Exception as inner_e:
                    logger.error(f"Failed to save partial report: {str(inner_e)}")
            
            raise e

    def process_key_files(self):
        """Process important files first to avoid duplicate analysis during directory exploration"""
        self.add_heading("Key File Analysis", level=1)
        
        # Define patterns for important files
        key_patterns = [
            {"pattern": "**/Cluster_Analysis.ipynb", "title": "Cluster Analysis", "method": self.process_notebook},
            {"pattern": "**/Feature_Importance.ipynb", "title": "Feature Importance Analysis", "method": self.process_notebook},
            {"pattern": "**/model_performance*.csv", "title": "Model Performance Analysis", "method": self.process_csv_file},
            {"pattern": "**/confusion_matrix*.png", "title": "Model Evaluation", "method": self.process_image_file},
            {"pattern": "**/feature_importance*.csv", "title": "Feature Importance Results", "method": self.process_csv_file},
            {"pattern": "**/feature_importance*.png", "title": "Feature Importance Visualization", "method": self.process_image_file},
        ]
        
        # Find and process key files
        key_files_found = False
        for pattern_info in key_patterns:
            pattern = pattern_info["pattern"]
            for root, dirs, files in os.walk(self.project_root):
                for file in glob.glob(os.path.join(root, os.path.basename(pattern))):
                    if file not in self.processed_files:
                        key_files_found = True
                        self.add_heading(f"{pattern_info['title']}: {os.path.basename(file)}", level=2)
                        pattern_info["method"](file)
                        self.processed_files.add(file)
        
        if not key_files_found:
            self.add_paragraph("No key analysis files were found in the project.")    

    
    def get_project_structure(self):
        """Get a text representation of the project structure for ChatGPT analysis"""
        structure = []
        
        def explore(dir_path, indent=0):
            """Helper function to explore directory structure"""
            try:
                items = sorted(os.listdir(dir_path))
                
                # Skip hidden items and __pycache__
                items = [item for item in items if not item.startswith('.') and item != '__pycache__']
                
                for item in items:
                    item_path = os.path.join(dir_path, item)
                    rel_path = os.path.relpath(item_path, self.project_root)
                    
                    if os.path.isdir(item_path):
                        structure.append("  " * indent + f"Directory: {rel_path}/")
                        # Limit recursion depth to avoid too much detail
                        if indent < 3:
                            explore(item_path, indent + 1)
                    else:
                        structure.append("  " * indent + f"File: {rel_path}")
            except Exception as e:
                logger.warning(f"Error exploring {dir_path} for structure: {str(e)}")
        
        explore(self.project_root)
        return "\n".join(structure)
    
    def process_specific_files(self):
        """Process specific important files that may need special attention"""
        self.add_heading("Key File Analysis", level=1)
        
        # Look for key files
        plot_file = os.path.join(self.project_root, "Code", "Plots.do")
        if os.path.exists(plot_file):
            self.add_heading("Plots.do Analysis", level=2)
            self.process_code_file(plot_file)
        
        # Check for specific notebooks
        cluster_nb = os.path.join(self.project_root, "Code", "Cluster_Analysis.ipynb")
        if os.path.exists(cluster_nb):
            self.add_heading("Cluster Analysis Notebook", level=2)
            self.process_notebook(cluster_nb)
        
        feature_nb = os.path.join(self.project_root, "Code", "Feature_Importance.ipynb")
        if os.path.exists(feature_nb):
            self.add_heading("Feature Importance Notebook", level=2)
            self.process_notebook(feature_nb)
        
        # Check for important results files
        for root, dirs, files in os.walk(os.path.join(self.project_root, "Output")):
            for file in files:
                if "model_performance" in file.lower() or "confusion_matrix" in file.lower() or "feature_importance" in file.lower():
                    file_path = os.path.join(root, file)
                    self.add_heading(f"Key Result: {file}", level=2)
                    
                    file_ext = os.path.splitext(file)[1].lower()
                    if file_ext == '.csv':
                        self.process_csv_file(file_path)
                    elif file_ext in ['.xlsx', '.xls']:
                        self.process_excel_file(file_path)
                    elif file_ext in ['.png', '.jpg', '.jpeg', '.gif']:
                        self.process_image_file(file_path)
    
    def generate_conclusion(self):
        """Generate a professional conclusion with context-aware insights"""
        self.add_heading("Conclusions and Recommendations", level=1)
        
        # Use the built project context for more meaningful conclusions
        analysis_types = list(self.project_context["identified_topics"])
        key_terms = list(self.project_context["key_terms"])
        
        # Create a more focused prompt based on what we've learned about the project
        if analysis_types:
            analysis_text = f"This project appears to involve {', '.join(analysis_types)}."
        else:
            analysis_text = "This project appears to involve data analysis with an unknown specific focus."
        
        if key_terms:
            terms_text = f"Key terms identified in the project include: {', '.join(key_terms)}."
        else:
            terms_text = "No specific domain terms were identified in the project."
        
        prompt = f"""
        Based on analysis of this data science project:
        
        {analysis_text}
        {terms_text}
        
        Generate a detailed and professional conclusion (around 350 words) that:
        1. Summarizes the apparent purpose and scope of the project
        2. Highlights the key analytical approaches used
        3. Suggests the potential implications of the findings
        4. Recommends possible next steps or areas for further investigation
        
        Make it professional, insightful, and action-oriented, as if it were the conclusion section of a formal executive report.
        Focus particularly on insights related to {', '.join(analysis_types) if analysis_types else 'data analysis'}.
        """
        
        conclusion = self._query_chatgpt(prompt)
        
        # Split into paragraphs for better readability
        paragraphs = conclusion.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                self.add_paragraph(paragraph)
        
        # Add recommendations section with context awareness
        self.add_heading("Recommendations for Next Steps", level=2)
        
        prompt = f"""
        Based on this project involving {', '.join(analysis_types) if analysis_types else 'data analysis'}
        and focusing on terms like {', '.join(key_terms[:5]) if key_terms else 'unknown terms'},
        provide a list of 5 concrete recommendations to improve or expand the analysis.
        
        Each recommendation should be specific, actionable, and relevant to the type of analysis performed.
        Focus on technical, methodological, and results presentation aspects.
        Make each recommendation 2-3 sentences, starting with an action verb.
        """
        
        recommendations = self._query_chatgpt(prompt)
        
        # Split and format as bullet points
        rec_list = recommendations.split('\n')
        for rec in rec_list:
            rec = rec.strip()
            if rec and not rec.isspace():
                # Remove numbers or dashes at the beginning if they exist
                rec = re.sub(r'^[\d\-\.\s]+', '', rec).strip()
                self.add_bullet_point(rec)


class PageNumCanvas(canvas.Canvas):
    """Canvas that adds page numbers to each page"""
    def __init__(self, *args, **kwargs):
        canvas.Canvas.__init__(self, *args, **kwargs)
        self.pages = []

    def showPage(self):
        self.pages.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        page_count = len(self.pages)
        for page in self.pages:
            self.__dict__.update(page)
            self.draw_page_number(page_count)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_page_number(self, page_count):
        # Skip page number on cover page
        if len(self.pages) == 1:
            return
            
        page = f"Page {len(self.pages)} of {page_count}"
        self.setFont("Helvetica", 9)
        self.setFillColor(colors.darkgrey)
        self.drawRightString(self._pagesize[0] - 50, 20, page)
        
        # Add a subtle line at the bottom
        self.setStrokeColor(colors.lightgrey)
        self.line(50, 30, self._pagesize[0] - 50, 30)

def main():
    """Main function with improved error handling and progress reporting"""
    try:
        # Setup logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[logging.FileHandler('report_generator.log'), logging.StreamHandler()]
        )
        
        # Get project root directory
        project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        
        print("Starting Project Report Generator...")
        print(f"Analyzing project in: {project_root}")
        
        # Initialize the report generator
        generator = ProjectReportGenerator(project_root)
        
        # Build project context first for better analysis
        print("Building project context...")
        generator.build_project_context()
        
        # Process the project with progress updates
        print("Generating report...")
        output_pdf = generator.process_project()
        
        print(f"Report generation complete! PDF saved to: {output_pdf}")
        print("You can now view the comprehensive project analysis.")
        return 0
        
    except Exception as e:
        logging.error(f"Error in main function: {str(e)}")
        print(f"Error generating report: {str(e)}")
        return 1

if __name__ == "__main__":
    main()